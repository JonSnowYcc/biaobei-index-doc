import tkinter as tk
from tkinter import filedialog
import os
import json
import re
import xlwings as xw
import threading
import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor  # 新增导入RGBColor类
from docx.enum.text import WD_COLOR_INDEX
from charset_normalizer import detect
import portalocker
import random



# 创建一个日志更新函数
def log_message(message):
    output_box.insert(tk.END, message + '\n')
    output_box.yview(tk.END)  # 滚动到最后一行

# 获取桌面路径
def get_desktop_path():
    return os.path.join(os.path.expanduser("~"), "Desktop")

# 提取文本中的标记内容
def extract_text_between_tags(text):
    pattern = r'〈(.*?)〉([^〈]*)'
    matches = re.findall(pattern, text)
    result = {f'〈{match[0]}〉': match[1].strip() for match in matches}
    return result

# 去除文本中的括号、拼音、英文字符等
def remove_bracketed_content(text, sheet=None, cell=None):
    bracket_pattern = r"\([^()]*\)|\[[^\[\]]*\]|\uFF08[^\uFF08\uFF09]*\uFF09|（[^（）]*）|（）|()"
    text_without_brackets = re.sub(bracket_pattern, '', text)
    text_without_pinyin_and_english = re.sub(r'[a-zA-Zǐǒáǎāǔóàōòìíī〉]+', '', text_without_brackets)
    text_without_commas_spaces = (text_without_pinyin_and_english.replace(',', '').replace('，', '')
                                  .replace(' ', '').replace('-', '').replace('“', '')
                                  .replace('”', '').replace('’', '').replace('）', '')
                                  .replace('(', '').replace('（', '').replace(')', '')
                                  .replace('"', '').replace('.', '').replace('。', '')
                                  .replace(';', '').replace('；', '').replace('。', ''))

    result = ""
    for char in text_without_commas_spaces:
        if char not in result:
            result += char

    return result

# 处理Excel文件
def process_excel(file_path):
    try:
        app = xw.App(visible=False)
        wb = app.books.open(file_path)
        sheet = wb.sheets[0]
        result = {}

        empty_row_count = 0  # 用于记录连续空行的计数

        for row in range(1, sheet.cells.last_cell.row + 1):
            row_empty = True
            for col in range(1, 6):
                cell = sheet.cells(row, col)
                if cell.value:
                    row_empty = False
                    break

            if row_empty:
                empty_row_count += 1
            else:
                empty_row_count = 0

            if empty_row_count >= 10:
                log_message(f"已遇到连续10行空行，退出处理。")
                break

            for col in range(1, 6):
                cell = sheet.cells(row, col)
                text = str(cell.value) if cell.value else ""
                if '〈' in text and '〉' in text:
                    extracted_texts = extract_text_between_tags(text)
                    for tag, extracted_text in extracted_texts.items():
                        clean_text = remove_bracketed_content(extracted_text, sheet, cell)
                        if not clean_text:
                            continue

                        part1 = sheet.cells(row, 2).value or ""
                        part2 = sheet.cells(row, 1).value or ""
                        part4 = find_value_of_first_no_border_cell_below(sheet, row, col)
                        for char in clean_text:
                            try:
                                value = f"{part1};{part2};{tag};{part4};"
                                value = re.sub(r'\s+', ' ', value)
                                if char in result:
                                    if len(result[char]) < 4:
                                        result[char].append(value)
                                else:
                                    result[char] = [value]

                                log_message(f"行数：{row} 处理字符 key {char} : {result[char]}")

                            except Exception as e:
                                log_message(f"处理字符 {char} 时发生错误: {e}")

        output_path = os.path.join(get_desktop_path(), "result.json")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=4)
        log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 处理完成，数据库文件已保存到: {output_path}")
        wb.close()
        app.quit()
    except Exception as e:
        log_message(f"处理Excel文件时发生错误: {str(e)}")

# 查找没有边框的单元格
def find_value_of_first_no_border_cell_below(sheet, row, col):
    try:
        for r in range(row - 1, 0, -1):
            cell = sheet.cells(r, col)
            borders = cell.api.Borders
            border_sides = [
                borders(xw.constants.BordersIndex.xlEdgeTop),
                borders(xw.constants.BordersIndex.xlEdgeBottom),
                borders(xw.constants.BordersIndex.xlEdgeLeft),
                borders(xw.constants.BordersIndex.xlEdgeRight),
            ]
            no_border_count = sum(1 for border in border_sides if border.LineStyle == -4142)
            if no_border_count >= 3:
                below_cell = sheet.cells(r + 1, col)
                below_value = below_cell.value
                return str(below_value).strip() if below_value is not None else ""

        return ""
    except Exception as e:
        log_message(f"发生错误: {e}")
        return ""


# 在button1_click和button2_click函数中相关位置替换原有的判断逻辑
# 例如在button2_click函数中类似这样替换
def button2_click():
    file_path = filedialog.askopenfilename(
        title="选择一个Excel文件",
        filetypes=[("Excel文件", "*.xlsx *.xls")]
    )
    if file_path:
        log_message(f"选择的文件路径为: {file_path}")
        threading.Thread(target=process_excel, args=(file_path,)).start()
    else:
        log_message("未选择任何文件")

import random

def process_txt_and_update_doc(txt_file_path, result_json_path, output_file_path):
    try:
        with open(txt_file_path, 'rb') as raw_file:
            raw_data = raw_file.read()
            result = detect(raw_data)
            encoding = result['encoding']

        with open(txt_file_path, 'r', encoding=encoding) as txt_file:
            txt_lines = txt_file.readlines()

        with open(result_json_path, 'r', encoding='utf-8') as json_file:
            result_data = json.load(json_file)

        current_document = Document()
        current_page_count = 0
        page_limit = 200
        base_output_path = output_file_path.rsplit('.', 1)[0]  # 获取不带后缀的基础输出路径

        # 设置默认段落格式
        style = current_document.styles['Normal']
        style.paragraph_format.space_after = Pt(0)  # 将段后间距设为0
        style.paragraph_format.line_spacing = 1.0  # 设置为单倍行距

        total_chars = sum(len(line) for line in txt_lines)
        processed_chars = 0
        next_progress = 2

        # 处理每个自然段
        for line_index, line in enumerate(txt_lines):
            # 在每个自然段开始前添加一条横线（除了第一段）
            if line_index > 0 and line.strip():
                paragraph = current_document.add_paragraph()
                # 设置段落格式
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                run = paragraph.add_run('_' * 60)
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            char_values = []
            first_part_char_map = {}
            second_part_char_map = {}

            for char in line:
                if char in result_data:
                    values = result_data[char]
                    char_values.append((char, values))
                    first_part_char_map[char] = set()
                    second_part_char_map[char] = set()

                    for value in values:
                        parts = value.split(';', 3)
                        if len(parts) >= 2:
                            first_part_char_map[char].add(parts[0])
                            second_part_char_map[char].add(parts[1])
                else:
                    char_values.append((char, None))

            first_part_colors = {}
            second_part_colors = {}
            processed_chars_set = set()

            for i, (char1, values1) in enumerate(char_values):
                if values1 is None or char1 in processed_chars_set:
                    continue

                for j, (char2, values2) in enumerate(char_values):
                    if values2 is None or char1 == char2 or char2 in processed_chars_set:
                        continue

                    # 扩展检索范围到前后两个字符
                    for offset in range(-2, 3):
                        if i + offset < 0 or i + offset >= len(sentence_chars):
                            continue
                        neighbor_char = sentence_chars[i + offset]
                        
                        # 比较前两项信息
                        if neighbor_char in first_part_char_map and char in first_part_char_map:
                            common_first = first_part_char_map[neighbor_char] & first_part_char_map[char]
                            if common_first:
                                for part in common_first:
                                    if part not in first_part_colors:
                                        first_part_colors[part] = RGBColor(
                                            random.randint(0, 255),
                                            random.randint(0, 255),
                                            random.randint(0, 255)
                                        )
                        
                        # 比较后两项信息
                        if neighbor_char in second_part_char_map and char in second_part_char_map:
                            common_second = second_part_char_map[neighbor_char] & second_part_char_map[char]
                            if common_second:
                                for part in common_second:
                                    if part not in second_part_colors:
                                        second_part_colors[part] = RGBColor(
                                            random.randint(0, 255),
                                            random.randint(0, 255),
                                            random.randint(0, 255)
                                        )
                        
                        # 两项完全匹配的情况
                        if common_first and common_second:
                            # 添加特殊标记符号3的逻辑
                            pass

                processed_chars_set.add(char1)

            for char, values in char_values:
                paragraph = current_document.add_paragraph()
                # 设置每个段落的格式
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

                processed_chars += 1
                progress = (processed_chars / total_chars) * 100

                if progress >= next_progress:
                    log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  Processing progress: {int(progress)}%")
                    next_progress += 2

                if values:
                    run_char = paragraph.add_run(char)
                    run_char.font.size = Pt(11)
                    run_char.font.name = '宋体'
                    run_char.font.color.rgb = RGBColor(255, 0, 0)
                    run_char.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    run_translation = paragraph.add_run("【")
                    run_translation.font.size = Pt(9)
                    run_translation.font.name = '宋体'
                    run_translation.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    for idx, value in enumerate(values):
                        parts = value.split(';', 3)
                        if len(parts) >= 2:
                            first_part = parts[0]
                            first_run = paragraph.add_run(first_part)
                            first_run.font.size = Pt(9)
                            first_run.font.bold = True
                            first_run.font.name = '宋体'
                            first_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            if first_part in first_part_colors:
                                first_run.font.color.rgb = first_part_colors[first_part]
                                first_run.font.underline = True
                                # 添加符号1标记（~）
                                symbol_run = paragraph.add_run(" ~")
                                symbol_run.font.size = Pt(9)
                                symbol_run.font.color.rgb = RGBColor(0, 128, 0)

                            semicolon1 = paragraph.add_run(";")
                            semicolon1.font.size = Pt(9)
                            semicolon1.font.bold = True
                            semicolon1.font.name = '宋体'
                            semicolon1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                            second_part = parts[1]
                            second_run = paragraph.add_run(second_part)
                            second_run.font.size = Pt(9)
                            second_run.font.bold = True
                            second_run.font.name = '宋体'
                            second_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            if second_part in second_part_colors:
                                second_run.font.color.rgb = second_part_colors[second_part]
                                # 添加符号2标记（*）
                                symbol_run = paragraph.add_run(" *")
                                symbol_run.font.size = Pt(9)
                                symbol_run.font.color.rgb = RGBColor(255, 0, 255)
                                
                                # 添加符号3标记（完全匹配⇒⇒）
                                if common_first and common_second:
                                    arrow_run = paragraph.add_run(" ⇒⇒")
                                    arrow_run.font.size = Pt(9)
                                    arrow_run.font.color.rgb = RGBColor(0, 0, 255)
                                second_run.font.underline = True

                            semicolon2 = paragraph.add_run("; ")
                            semicolon2.font.size = Pt(9)
                            semicolon2.font.bold = True
                            semicolon2.font.name = '宋体'
                            semicolon2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                            if len(parts) > 2:
                                remaining_text = parts[2]
                                remaining_run = paragraph.add_run(remaining_text)
                                remaining_run.font.size = Pt(9)
                                remaining_run.font.name = '宋体'
                                remaining_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                        if idx < len(values) - 1:
                            separator = paragraph.add_run(" | ")
                            separator.font.size = Pt(9)
                            separator.font.name = '宋体'
                            separator.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    run_translation_end = paragraph.add_run("】")
                    run_translation_end.font.size = Pt(9)
                    run_translation_end.font.name = '宋体'
                    run_translation_end.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                else:
                    run = paragraph.add_run(char)
                    run.font.size = Pt(11)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 检查当前文档页数是否达到限制，达到则新建文档
            if current_page_count >= page_limit:
                current_document.save(f"{base_output_path}_{current_page_count // page_limit + 1}.doc")
                current_document = Document()
                current_page_count = 0
                # 重新设置新文档的默认段落格式
                style = current_document.styles['Normal']
                style.paragraph_format.space_after = Pt(0)
                style.paragraph_format.line_spacing = 1.0

        # 保存最后一个未满200页的文档
        current_document.save(f"{base_output_path}_{current_page_count // page_limit + 1}.doc")

        log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 新文件已生成并保存")
    except Exception as e:
        log_message(f"处理TXT文件时发生错误: {str(e)}")

# 修改按钮点击函数，确保传递了三个参数
def button1_click():
    txt_file_path = filedialog.askopenfilename(
        title="选择一个txt文件",
        filetypes=[("Text文件", "*.txt")]
    )
    if txt_file_path:
        log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 选择的txt文件路径为: {txt_file_path}")
        log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 检索文件中，请等待....")
        result_json_path = os.path.join(get_desktop_path(), "result.json")
        output_file_path = os.path.join(get_desktop_path(), "updated_result.doc")
        threading.Thread(target=process_txt_and_update_doc, args=(txt_file_path, result_json_path, output_file_path)).start()
    else:
        log_message("未选择任何txt文件")

# 创建GUI窗口
root = tk.Tk()
root.title("Excel处理工具")
root.geometry("1200x400")

# 添加Text组件用于输出日志
output_box = tk.Text(root, height=15, width=600)
output_box.pack(pady=20)

button2 = tk.Button(root, text="处理Excel", command=button2_click)
button2.pack(pady=10)

button1 = tk.Button(root, text="处理TXT", command=button1_click)
button1.pack(pady=10)

root.mainloop()

