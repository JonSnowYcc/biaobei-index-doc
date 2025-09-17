import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
import os
import json
import re
import xlwings as xw
import threading
import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor  # 新增导入RGBColor类
from docx.enum.text import WD_COLOR_INDEX
from charset_normalizer import detect
import portalocker
import random



# 创建一个日志更新函数
def log_message(message):
    output_box.insert(tk.END, message + '\n')
    output_box.yview(tk.END)  # 滚动到最后一行

# 获取桌面路径
def get_desktop_path():
    return os.path.join(os.path.expanduser("~"), "Desktop")

# 提取文本中的标记内容
def extract_text_between_tags(text):
    pattern = r'〈(.*?)〉([^〈]*)'
    matches = re.findall(pattern, text)
    result = {f'〈{match[0]}〉': match[1].strip() for match in matches}
    return result

# 提取第二部分（部名），如"元部開四[ian]"中的"元部"
def extract_second_part(second_part_full):
    """从'元部開四[ian]'中提取'元部'"""
    if not second_part_full:
        return ""
    # 查找"部"字的位置，提取到"部"为止
    bu_index = second_part_full.find('部')
    if bu_index != -1:
        return second_part_full[:bu_index + 1]  # 包含"部"字
    return second_part_full


# 去除文本中的括号、拼音、英文字符等
def remove_bracketed_content(text, sheet=None, cell=None):
    bracket_pattern = r"\([^()]*\)|\[[^\[\]]*\]|\uFF08[^\uFF08\uFF09]*\uFF09|（[^（）]*）|（）|()"
    text_without_brackets = re.sub(bracket_pattern, '', text)
    text_without_pinyin_and_english = re.sub(r'[a-zA-Zǐǒáǎāǔóàōòìíī〉]+', '', text_without_brackets)
    text_without_commas_spaces = (text_without_pinyin_and_english.replace(',', '').replace('，', '')
                                  .replace(' ', '').replace('-', '').replace('"', '')
                                  .replace('"', '').replace('"', '').replace('"', '').replace('"', '')
                                  .replace('.', '').replace('。', '')
                                  .replace(';', '').replace('；', '').replace('。', ''))

    result = ""
    for char in text_without_commas_spaces:
        if char not in result:
            result += char

    return result

# 处理Excel文件
def process_excel(file_path):
    try:
        app = xw.App(visible=False)
        wb = app.books.open(file_path)
        sheet = wb.sheets[0]
        result = {}

        empty_row_count = 0  # 用于记录连续空行的计数

        total_rows = sheet.cells.last_cell.row
        processed_rows = 0
        # 初始化进度
        def _set_progress_excel(pct):
            try:
                root.after(0, lambda: progress_var.set(max(0, min(100, pct))))
                root.after(0, lambda: progress_label_var.set(f"Excel处理进度：{max(0, min(100, pct)):.1f}%"))
            except Exception:
                pass

        _set_progress_excel(0)

        for row in range(1, total_rows + 1):
            row_empty = True
            for col in range(1, 6):
                cell = sheet.cells(row, col)
                if cell.value:
                    row_empty = False
                    break

            if row_empty:
                empty_row_count += 1
            else:
                empty_row_count = 0

            if empty_row_count >= 10:
                log_message(f"已遇到连续10行空行，退出处理。")
                break

            for col in range(1, 6):
                cell = sheet.cells(row, col)
                text = str(cell.value) if cell.value else ""
                if '〈' in text and '〉' in text:
                    extracted_texts = extract_text_between_tags(text)
                    for tag, extracted_text in extracted_texts.items():
                        clean_text = remove_bracketed_content(extracted_text, sheet, cell)
                        if not clean_text:
                            continue

                        part1 = sheet.cells(row, 2).value or ""
                        part2 = sheet.cells(row, 1).value or ""
                        part4 = find_value_of_first_no_border_cell_below(sheet, row, col)
                        for char in clean_text:
                            try:
                                value = f"{part1};{part2};{tag};{part4};"
                                value = re.sub(r'\s+', ' ', value)
                                if char in result:
                                    if len(result[char]) < 4:
                                        result[char].append(value)
                                else:
                                    result[char] = [value]

                                log_message(f"行数：{row} 处理字符 key {char} : {result[char]}")

                            except Exception as e:
                                log_message(f"处理字符 {char} 时发生错误: {e}")

            processed_rows += 1
            if total_rows > 0:
                _set_progress_excel(processed_rows / total_rows * 100)

        output_path = os.path.join(get_desktop_path(), "result.json")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=4)
        log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 处理完成，数据库文件已保存到: {output_path}")
        _set_progress_excel(100)
        wb.close()
        app.quit()
    except Exception as e:
        log_message(f"处理Excel文件时发生错误: {str(e)}")
        try:
            root.after(0, lambda: progress_label_var.set("Excel处理出错"))
        except Exception:
            pass

# 查找没有边框的单元格
def find_value_of_first_no_border_cell_below(sheet, row, col):
    try:
        for r in range(row - 1, 0, -1):
            cell = sheet.cells(r, col)
            borders = cell.api.Borders
            border_sides = [
                borders(xw.constants.BordersIndex.xlEdgeTop),
                borders(xw.constants.BordersIndex.xlEdgeBottom),
                borders(xw.constants.BordersIndex.xlEdgeLeft),
                borders(xw.constants.BordersIndex.xlEdgeRight),
            ]
            no_border_count = sum(1 for border in border_sides if border.LineStyle == -4142)
            if no_border_count >= 3:
                below_cell = sheet.cells(r + 1, col)
                below_value = below_cell.value
                return str(below_value).strip() if below_value is not None else ""

        return ""
    except Exception as e:
        log_message(f"发生错误: {e}")
        return ""


# 在button1_click和button2_click函数中相关位置替换原有的判断逻辑
# 例如在button2_click函数中类似这样替换
def button2_click():
    file_path = filedialog.askopenfilename(
        title="选择一个Excel文件",
        filetypes=[("Excel文件", "*.xlsx *.xls")]
    )
    if file_path:
        log_message(f"选择的文件路径为: {file_path}")
        threading.Thread(target=process_excel, args=(file_path,)).start()
    else:
        log_message("未选择任何文件")

import random

def process_txt_and_update_doc(txt_file_path, result_json_path, output_file_path):
    try:
        with open(txt_file_path, 'rb') as raw_file:
            raw_data = raw_file.read()
            result = detect(raw_data)
            encoding = result['encoding']

        with open(txt_file_path, 'r', encoding=encoding) as txt_file:
            txt_content = txt_file.read()

        with open(result_json_path, 'r', encoding='utf-8') as json_file:
            result_data = json.load(json_file)

        # 先将文本按换行符分割成段落，然后按标点符号分割成句子
        paragraphs = txt_content.split('\n')
        sentences = []  # 存储句子内容和分隔信息

        # 识别标题段落：无任何标点符号的非空行
        punctuation_marks = '。！？；，：'
        title_paragraph_indices = set()
        for para_idx, paragraph in enumerate(paragraphs):
            if not paragraph.strip():
                continue
            if not any(ch in punctuation_marks for ch in paragraph):
                title_paragraph_indices.add(para_idx)

        # 构建句子列表，并标记是否诗句（位于标题之后至下一标题之间）
        in_poem_section = False
        current_title_text = None
        titles_ordered = []
        for para_idx, paragraph in enumerate(paragraphs):
            if not paragraph.strip():  # 跳过空段落
                continue

            is_title_paragraph = para_idx in title_paragraph_indices
            # 标题段落：整段为一个标题句，无标点
            if is_title_paragraph:
                title_text = paragraph.strip()
                sentences.append({
                    'content': title_text,
                    'is_title': True,
                    'paragraph_index': para_idx,
                    'is_poem_section': False
                })
                # 自此之后进入诗句区，直到遇到下一个标题
                in_poem_section = True
                current_title_text = title_text
                titles_ordered.append(title_text)
                continue

            # 非标题段落：在段落内按标点符号分割成句子
            current_sentence = ''
            for i, char in enumerate(paragraph):
                current_sentence += char
                if char in punctuation_marks:
                    if current_sentence.strip():
                        sentences.append({
                            'content': current_sentence,
                            'is_title': False,
                            'paragraph_index': para_idx,
                            'is_poem_section': in_poem_section,
                            'title_text': current_title_text
                        })
                    current_sentence = ''

            # 处理段落中最后一个句子（如果没有以标点符号结尾）
            if current_sentence.strip():
                sentences.append({
                    'content': current_sentence,
                    'is_title': False,
                    'paragraph_index': para_idx,
                    'is_poem_section': in_poem_section,
                    'title_text': current_title_text
                })

        current_document = Document()
        current_page_count = 0
        page_limit = 200
        
        # 生成动态文件名格式：字典检索结果_20250913_101209_1
        current_time = datetime.datetime.now()
        timestamp = current_time.strftime("%Y%m%d_%H%M%S")
        base_output_path = f"字典检索结果_{timestamp}"

        # 设置默认段落格式
        style = current_document.styles['Normal']
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        # 统计变量
        total_sentences = 0
        all_sentence_stats = []  # 存储所有句子的统计信息
        poem_total_sentences = 0
        poem_marked_sentences = 0
        title_to_poem_counts = {}
        title_to_marked_counts = {}

        # 先统计总句子数（不包括标题）
        for sentence_info in sentences:
            if not sentence_info['is_title']:
                total_sentences += 1

        # 全局总览输出已取消（按标题统计见文末）

        # 进度初始化
        total_sentences_for_progress = max(1, len(sentences))
        processed_sentences_for_progress = 0
        def _set_progress_txt(pct):
            try:
                root.after(0, lambda: progress_var.set(max(0, min(100, pct))))
                root.after(0, lambda: progress_label_var.set(f"TXT生成进度：{max(0, min(100, pct)):.1f}%"))
            except Exception:
                pass

        _set_progress_txt(0)

        # 处理每个句子
        for sentence_info in sentences:
            sentence = sentence_info['content']
            is_title = sentence_info['is_title']
            is_poem_section = sentence_info.get('is_poem_section', False)
            title_text = sentence_info.get('title_text')
            chars = list(sentence)
            marks = [[] for _ in range(len(chars))]  # 用于存储每个字符的标记
            
            
            # 初始化颜色映射字典（每个句子重新初始化）
            first_part_colors = {}
            second_part_colors = {}
            
            # 定义匹配类型的背景色（通过Word的highlight_color设置）
            # ① 浅绿色 - 前半段匹配
            # ② 黄色 - 后半段匹配  
            # ③ 粉色 - 都匹配

            # 在每个自然段开始前添加一条横线（除了第一段）
            if len(sentences) > 0:
                paragraph = current_document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                run = paragraph.add_run('_' * 60)
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 收集句子内所有字符的第一项和第二项信息
            sentence_first_parts = {}  # 第一项 -> 字符列表（不是索引）
            sentence_second_parts = {}  # 第二项（部名） -> 字符列表（不是索引）
            
            for i in range(len(chars)):
                if chars[i] == '。' or chars[i] not in result_data:
                    continue
                    
                current_info = result_data[chars[i]]
                
                for current_value in current_info:
                    current_parts = current_value.split(';', 3)
                    if len(current_parts) >= 2:
                        first_part = current_parts[0]
                        second_part_full = current_parts[1]
                        second_part = extract_second_part(second_part_full)  # 提取部名
                        
                        # 记录第一项 - 按字符记录，不是按索引
                        if first_part not in sentence_first_parts:
                            sentence_first_parts[first_part] = []
                        if chars[i] not in sentence_first_parts[first_part]:
                            sentence_first_parts[first_part].append(chars[i])
                        
                        # 记录第二项（部名） - 按字符记录，不是按索引
                        if second_part and second_part not in sentence_second_parts:
                            sentence_second_parts[second_part] = []
                        if second_part and chars[i] not in sentence_second_parts[second_part]:
                            sentence_second_parts[second_part].append(chars[i])
            
            # 创建第一部分的匹配组 - 将相似的第一部分归为一组
            first_part_groups = {}  # 组名 -> 字符列表
            for first_part, char_list in sentence_first_parts.items():
                # 提取第一部分的主要部分（去掉数字和特殊字符）
                import re
                main_part = re.sub(r'[0-9lʧȶøXbʣp\'k\']', '', first_part)  # 去掉数字和特殊字符
                if not main_part:  # 如果去掉特殊字符后为空，使用原字符串
                    main_part = first_part
                
                if main_part not in first_part_groups:
                    first_part_groups[main_part] = []
                first_part_groups[main_part].extend(char_list)
                # 去重
                first_part_groups[main_part] = list(set(first_part_groups[main_part]))
            
            # 为有多个字符匹配的第一项和第二项添加标记
            for group_name, char_list in first_part_groups.items():
                if len(char_list) > 1:  # 只有多个不同字符匹配时才标记
                    # 为所有匹配的字符添加标记
                    for char in char_list:
                        for i in range(len(chars)):
                            if chars[i] == char:
                                if '①' not in marks[i]: marks[i].append('①')
                    
                    # 为这个组的所有第一部分分配相同的颜色
                    group_color = RGBColor(
                        random.randint(0, 255),
                        random.randint(0, 255),
                        random.randint(0, 255)
                    )
                    for first_part in sentence_first_parts:
                        # 检查这个第一部分是否属于当前组
                        main_part = re.sub(r'[0-9lʧȶøXbʣp\'k\']', '', first_part)
                        if not main_part:
                            main_part = first_part
                        if main_part == group_name:
                            first_part_colors[first_part] = group_color
            
            for second_part, char_list in sentence_second_parts.items():
                if len(char_list) > 1:  # 只有多个不同字符匹配时才标记
                    # 为所有匹配的字符添加标记
                    for char in char_list:
                        for i in range(len(chars)):
                            if chars[i] == char:
                                if '②' not in marks[i]: marks[i].append('②')
                    # 只有匹配上的第二部分才分配颜色
                    if second_part not in second_part_colors:
                        second_part_colors[second_part] = RGBColor(
                            random.randint(0, 255),
                            random.randint(0, 255),
                            random.randint(0, 255)
                        )
            
            # 收集所有字符的1和2部分组合，用于判断是否应该标记③
            char_combinations = {}  # 字符 -> 1和2部分组合列表
            combination_to_chars = {}  # 1和2部分组合 -> 字符列表
            
            for i in range(len(chars)):
                if chars[i] == '。' or chars[i] not in result_data:
                    continue
                    
                current_info = result_data[chars[i]]
                char_combinations[chars[i]] = []
                
                for current_value in current_info:
                    current_parts = current_value.split(';', 3)
                    if len(current_parts) >= 2:
                        first_part = current_parts[0]
                        second_part_full = current_parts[1]
                        second_part = extract_second_part(second_part_full)  # 提取部名
                        
                        # 创建1和2部分的组合
                        combination = f"{first_part}|{second_part}"
                        char_combinations[chars[i]].append(combination)
                        
                        # 记录这个组合对应的字符
                        if combination not in combination_to_chars:
                            combination_to_chars[combination] = []
                        if chars[i] not in combination_to_chars[combination]:
                            combination_to_chars[combination].append(chars[i])
            
            # 检查哪些字符既有第一项匹配又有第二项匹配，并且1和2组合与其他字符相同
            for i in range(len(chars)):
                if chars[i] == '。' or chars[i] not in result_data:
                    continue
                    
                has_first_match = '①' in marks[i]
                has_second_match = '②' in marks[i]
                
                # 检查是否有相同的1和2组合
                has_same_combination = False
                if chars[i] in char_combinations:
                    for combination in char_combinations[chars[i]]:
                        if combination in combination_to_chars and len(combination_to_chars[combination]) > 1:
                            has_same_combination = True
                            break
                
                if has_first_match and has_second_match and has_same_combination:
                    if '③' not in marks[i]: marks[i].append('③')


            # 显示每个字符及其标记
            for i, char in enumerate(chars):
                # 为每个字符创建新的段落
                paragraph = current_document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

                run_char = paragraph.add_run(char)
                run_char.font.size = Pt(11)
                run_char.font.name = '宋体'
                # 如果是句号，不设置红色
                if char != '。':
                    run_char.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # 根据标记类型设置背景色
                    if marks[i]:
                        if '③' in marks[i]:  # 都匹配 - 浅紫色
                            run_char.font.highlight_color = WD_COLOR_INDEX.TURQUOISE  # 使用青绿色替代，更清晰
                        elif '①' in marks[i]:  # 前半段匹配 - 浅绿色
                            run_char.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        elif '②' in marks[i]:  # 后半段匹配 - 黄色
                            run_char.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            
                run_char.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                if char != '。' and char in result_data:
                    # 添加标记符号
                    if marks[i]:
                        mark_text = ''.join(marks[i])
                        mark_run = paragraph.add_run(mark_text)
                        mark_run.font.size = Pt(9)
                        mark_run.font.name = '宋体'
                        mark_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    # 添加原有的注释信息
                    run_translation = paragraph.add_run("【")
                    run_translation.font.size = Pt(9)
                    run_translation.font.name = '宋体'
                    run_translation.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    for idx, value in enumerate(result_data[char]):
                        parts = value.split(';', 3)
                        if len(parts) >= 2:
                            first_part = parts[0]
                            first_run = paragraph.add_run(first_part)
                            first_run.font.size = Pt(9)
                            first_run.font.bold = True
                            first_run.font.name = '宋体'
                            first_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            if first_part in first_part_colors:
                                first_run.font.color.rgb = first_part_colors[first_part]
                                first_run.font.underline = True

                            semicolon1 = paragraph.add_run(";")
                            semicolon1.font.size = Pt(9)
                            semicolon1.font.bold = True
                            semicolon1.font.name = '宋体'
                            semicolon1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                            second_part = parts[1]
                            # 提取部名用于匹配判断
                            bu_name = extract_second_part(second_part)
                            
                            # 分段显示第二部分，给部名加下划线和字体颜色
                            if bu_name and bu_name in second_part_colors:
                                # 使用正则表达式找到部名的位置，确保只匹配第一个"部"字
                                import re
                                pattern = r'(.+?)(部)(.*)'
                                match = re.match(pattern, second_part)
                                if match:
                                    before_bu = match.group(1)  # 部名前的部分（如"元"）
                                    bu_char = match.group(2)    # "部"字
                                    after_bu = match.group(3)   # 部名后的部分
                                    
                                    # 部名前的部分 - 加下划线和字体颜色（像第一部分一样）
                                    if before_bu:
                                        before_run = paragraph.add_run(before_bu)
                                        before_run.font.size = Pt(9)
                                        before_run.font.bold = True
                                        before_run.font.name = '宋体'
                                        before_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                        # 加下划线和字体颜色
                                        before_run.font.color.rgb = second_part_colors[bu_name]
                                        before_run.font.underline = True
                                    
                                    # 部字 - 加下划线和字体颜色
                                    bu_run = paragraph.add_run(bu_char)  # 只显示"部"字
                                    bu_run.font.size = Pt(9)
                                    bu_run.font.bold = True
                                    bu_run.font.name = '宋体'
                                    bu_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                    # 加下划线和字体颜色
                                    bu_run.font.color.rgb = second_part_colors[bu_name]
                                    bu_run.font.underline = True
                                    
                                    # 部名后的部分
                                    if after_bu:
                                        after_run = paragraph.add_run(after_bu)
                                        after_run.font.size = Pt(9)
                                        after_run.font.bold = True
                                        after_run.font.name = '宋体'
                                        after_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                else:
                                    # 没找到匹配，正常显示
                                    second_run = paragraph.add_run(second_part)
                                    second_run.font.size = Pt(9)
                                    second_run.font.bold = True
                                    second_run.font.name = '宋体'
                                    second_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            else:
                                # 没有匹配的部名，正常显示
                                second_run = paragraph.add_run(second_part)
                                second_run.font.size = Pt(9)
                                second_run.font.bold = True
                                second_run.font.name = '宋体'
                                second_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                            semicolon2 = paragraph.add_run("; ")  # 恢复这行
                            semicolon2.font.size = Pt(9)
                            semicolon2.font.bold = True
                            semicolon2.font.name = '宋体'
                            semicolon2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                            if len(parts) > 2:
                                remaining_text = parts[2]
                                remaining_run = paragraph.add_run(remaining_text)
                                remaining_run.font.size = Pt(9)
                                remaining_run.font.name = '宋体'
                                remaining_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                        if idx < len(result_data[char]) - 1:
                            separator = paragraph.add_run(" | ")
                            separator.font.size = Pt(9)
                            separator.font.name = '宋体'
                            separator.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    run_translation_end = paragraph.add_run("】")
                    run_translation_end.font.size = Pt(9)
                    run_translation_end.font.name = '宋体'
                    run_translation_end.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 如果是标题，添加"（标题不统计）"标记
            if is_title:
                title_paragraph = current_document.add_paragraph()
                title_paragraph.paragraph_format.space_after = Pt(0)
                title_paragraph.paragraph_format.line_spacing = 1.0
                title_run = title_paragraph.add_run("（标题不统计）")
                title_run.font.size = Pt(9)
                title_run.font.name = '宋体'
                title_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
                title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                # 统计①、②、③的数量和触发的字典内容
                mark1_contents = {}  # 第一项内容 -> 触发次数
                mark2_contents = {}  # 第二项内容（部名） -> 触发次数
                mark3_contents = {}  # 第三项内容（1和2组合） -> 触发次数
                
                # 用于记录每个字符已经统计过的部分，避免重复计算
                char_mark1_parts = {}  # 字符 -> 已统计的第一项集合
                char_mark2_parts = {}  # 字符 -> 已统计的第二项集合
                char_mark3_parts = {}  # 字符 -> 已统计的组合集合
                
                for i, char in enumerate(chars):
                    if char == '。' or char not in result_data:
                        continue
                    
                    # 初始化字符的统计集合
                    if char not in char_mark1_parts:
                        char_mark1_parts[char] = set()
                    if char not in char_mark2_parts:
                        char_mark2_parts[char] = set()
                    if char not in char_mark3_parts:
                        char_mark3_parts[char] = set()
                    
                    if '①' in marks[i]:
                        # 收集①的内容（第一项）- 每个字符的相同第一项只计算一次
                        if char in result_data:
                            for value in result_data[char]:
                                parts = value.split(';', 3)
                                if len(parts) >= 1:
                                    first_part = parts[0]
                                    if first_part not in char_mark1_parts[char]:
                                        char_mark1_parts[char].add(first_part)
                                        if first_part not in mark1_contents:
                                            mark1_contents[first_part] = 0
                                        mark1_contents[first_part] += 1
                    
                    if '②' in marks[i]:
                        # 收集②的内容（第二项，部名）- 每个字符的相同第二项只计算一次
                        if char in result_data:
                            for value in result_data[char]:
                                parts = value.split(';', 3)
                                if len(parts) >= 2:
                                    second_part = extract_second_part(parts[1])
                                    if second_part and second_part not in char_mark2_parts[char]:
                                        char_mark2_parts[char].add(second_part)
                                        if second_part not in mark2_contents:
                                            mark2_contents[second_part] = 0
                                        mark2_contents[second_part] += 1
                    
                    if '③' in marks[i]:
                        # 收集③的内容（1和2的组合）- 每个字符的相同组合只计算一次
                        if char in result_data:
                            for value in result_data[char]:
                                parts = value.split(';', 3)
                                if len(parts) >= 2:
                                    first_part = parts[0]
                                    second_part = extract_second_part(parts[1])
                                    if second_part:
                                        combination = f"{first_part}|{second_part}"
                                        if combination not in char_mark3_parts[char]:
                                            char_mark3_parts[char].add(combination)
                                            if combination not in mark3_contents:
                                                mark3_contents[combination] = 0
                                            mark3_contents[combination] += 1
                
                # 收集统计信息到列表
                sentence_stats = {
                    'sentence': sentence,
                    'mark1_contents': mark1_contents,
                    'mark2_contents': mark2_contents,
                    'mark3_contents': mark3_contents,
                    'is_poem_section': is_poem_section,
                    'title_text': title_text
                }
                all_sentence_stats.append(sentence_stats)
                
                # 输出统计信息
                if mark1_contents or mark2_contents or mark3_contents:
                    stats_paragraph = current_document.add_paragraph()
                    stats_paragraph.paragraph_format.space_after = Pt(0)
                    stats_paragraph.paragraph_format.line_spacing = 1.0
                    
                    # 统计标题
                    stats_title = stats_paragraph.add_run("【统计信息】")
                    stats_title.font.size = Pt(9)
                    stats_title.font.name = '宋体'
                    stats_title.font.bold = True
                    stats_title.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                    stats_title.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    
                    # ①统计（只显示触发次数大于1的）
                    mark1_filtered = {content: count for content, count in mark1_contents.items() if count > 1}
                    if mark1_filtered:
                        mark1_line = current_document.add_paragraph()
                        mark1_line.paragraph_format.space_after = Pt(0)
                        mark1_line.paragraph_format.line_spacing = 1.0
                        mark1_items = [f"{content}({count})" for content, count in sorted(mark1_filtered.items())]
                        mark1_text = f"①触发内容: {', '.join(mark1_items)}"
                        mark1_run = mark1_line.add_run(mark1_text)
                        mark1_run.font.size = Pt(9)
                        mark1_run.font.name = '宋体'
                        mark1_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                        mark1_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    
                    # ②统计（只显示触发次数大于1的）
                    mark2_filtered = {content: count for content, count in mark2_contents.items() if count > 1}
                    if mark2_filtered:
                        mark2_line = current_document.add_paragraph()
                        mark2_line.paragraph_format.space_after = Pt(0)
                        mark2_line.paragraph_format.line_spacing = 1.0
                        mark2_items = [f"{content}({count})" for content, count in sorted(mark2_filtered.items())]
                        mark2_text = f"②触发内容: {', '.join(mark2_items)}"
                        mark2_run = mark2_line.add_run(mark2_text)
                        mark2_run.font.size = Pt(9)
                        mark2_run.font.name = '宋体'
                        mark2_run.font.color.rgb = RGBColor(255, 165, 0)  # 橙色
                        mark2_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    
                    # ③统计（只显示触发次数大于1的）
                    mark3_filtered = {content: count for content, count in mark3_contents.items() if count > 1}
                    if mark3_filtered:
                        mark3_line = current_document.add_paragraph()
                        mark3_line.paragraph_format.space_after = Pt(0)
                        mark3_line.paragraph_format.line_spacing = 1.0
                        mark3_items = [f"{content}({count})" for content, count in sorted(mark3_filtered.items())]
                        mark3_text = f"③触发内容: {', '.join(mark3_items)}"
                        mark3_run = mark3_line.add_run(mark3_text)
                        mark3_run.font.size = Pt(9)
                        mark3_run.font.name = '宋体'
                        mark3_run.font.color.rgb = RGBColor(128, 0, 128)  # 紫色
                        mark3_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 诗句计数（使用已有统计结果，不另做判定）
                if is_poem_section:
                    poem_total_sentences += 1
                    # 初始化标题聚合容器
                    if title_text:
                        if title_text not in title_to_poem_counts:
                            title_to_poem_counts[title_text] = 0
                        if title_text not in title_to_marked_counts:
                            title_to_marked_counts[title_text] = 0
                        title_to_poem_counts[title_text] += 1
                    if mark1_contents or mark2_contents or mark3_contents:
                        poem_marked_sentences += 1
                        if title_text:
                            title_to_marked_counts[title_text] += 1

            # 检查页数限制逻辑保持不变
            if current_page_count >= page_limit:
                current_document.save(f"{base_output_path}_{current_page_count // page_limit + 1}.doc")
                current_document = Document()
                current_page_count = 0
                style = current_document.styles['Normal']
                style.paragraph_format.space_after = Pt(0)
                style.paragraph_format.line_spacing = 1.0

            # 更新进度
            processed_sentences_for_progress += 1
            _set_progress_txt(processed_sentences_for_progress / total_sentences_for_progress * 100)

        # 在保存前仅输出“按标题诗句统计”与标题下的具体统计
        if all_sentence_stats:
            if titles_ordered:
                per_title_header = current_document.add_paragraph()
                per_title_header.paragraph_format.space_after = Pt(6)
                per_title_header.paragraph_format.line_spacing = 1.0
                per_title_header_run = per_title_header.add_run("按标题诗句统计")
                per_title_header_run.font.size = Pt(12)
                per_title_header_run.font.name = '宋体'
                per_title_header_run.font.bold = True
                per_title_header_run.font.color.rgb = RGBColor(0, 0, 0)
                per_title_header_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                for t in titles_ordered:
                    total_c = title_to_poem_counts.get(t, 0)
                    marked_c = title_to_marked_counts.get(t, 0)
                    ratio_t = (marked_c / total_c * 100) if total_c > 0 else 0.0
                    line = current_document.add_paragraph()
                    line.paragraph_format.space_after = Pt(3)
                    line.paragraph_format.line_spacing = 1.0
                    text = f"《{t}》：诗句{total_c}句；有标记{marked_c}句（{ratio_t:.1f}%）"
                    run = line.add_run(text)
                    run.font.size = Pt(11)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    # 小标题：该标题下具体统计
                    detail_header = current_document.add_paragraph()
                    detail_header.paragraph_format.space_after = Pt(3)
                    detail_header.paragraph_format.line_spacing = 1.0
                    detail_header_run = detail_header.add_run("—— 具体统计 ——")
                    detail_header_run.font.size = Pt(10)
                    detail_header_run.font.name = '宋体'
                    detail_header_run.font.color.rgb = RGBColor(128, 128, 128)
                    detail_header_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    # 遍历该标题下所有诗句，输出与已展示口径一致的行
                    for stats in all_sentence_stats:
                        if not stats.get('is_poem_section'):
                            continue
                        if stats.get('title_text') != t:
                            continue

                        sentence = stats['sentence']
                        mark1_contents = stats['mark1_contents']
                        mark2_contents = stats['mark2_contents']
                        mark3_contents = stats['mark3_contents']

                        mark1_filtered = {content: count for content, count in mark1_contents.items() if count > 1}
                        mark2_filtered = {content: count for content, count in mark2_contents.items() if count > 1}
                        mark3_filtered = {content: count for content, count in mark3_contents.items() if count > 1}

                        if not (mark1_filtered or mark2_filtered or mark3_filtered):
                            continue

                        line_parts = [f"{sentence}"]
                        if mark1_filtered:
                            mark1_items = [f"{content}({count})" for content, count in sorted(mark1_filtered.items())]
                            line_parts.append(f"①{', '.join(mark1_items)}")
                        if mark2_filtered:
                            mark2_items = [f"{content}({count})" for content, count in sorted(mark2_filtered.items())]
                            line_parts.append(f"②{', '.join(mark2_items)}")
                        if mark3_filtered:
                            mark3_items = [f"{content}({count})" for content, count in sorted(mark3_filtered.items())]
                            line_parts.append(f"③{', '.join(mark3_items)}")

                        detail_line = current_document.add_paragraph()
                        detail_line.paragraph_format.space_after = Pt(2)
                        detail_line.paragraph_format.line_spacing = 1.0
                        detail_text = "　　".join(line_parts)
                        detail_run = detail_line.add_run(detail_text)
                        detail_run.font.size = Pt(10)
                        detail_run.font.name = '宋体'
                        detail_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 保存最后一个文档
        current_document.save(f"{base_output_path}_{current_page_count // page_limit + 1}.doc")
        log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 新文件已生成并保存")
        _set_progress_txt(100)

    except Exception as e:
        log_message(f"处理TXT文件时发生错误: {str(e)}")
        try:
            root.after(0, lambda: progress_label_var.set("TXT处理出错"))
        except Exception:
            pass

        # 修改按钮点击函数，确保传递了三个参数
def button1_click():
    txt_file_path = filedialog.askopenfilename(
        title="选择一个txt文件",
        filetypes=[("Text文件", "*.txt")]
    )
    if txt_file_path:
        log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 选择的txt文件路径为: {txt_file_path}")
        log_message(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 检索文件中，请等待....")
        result_json_path = os.path.join(get_desktop_path(), "result.json")
        # 生成动态文件名格式：字典检索结果_20250913_101209_1
        current_time = datetime.datetime.now()
        timestamp = current_time.strftime("%Y%m%d_%H%M%S")
        output_file_path = os.path.join(get_desktop_path(), f"字典检索结果_{timestamp}.doc")
        threading.Thread(target=process_txt_and_update_doc, args=(txt_file_path, result_json_path, output_file_path)).start()
    else:
        log_message("未选择任何txt文件")

# 创建GUI窗口
root = tk.Tk()
root.title("Excel处理工具")
root.geometry("1200x400")

# 添加Text组件用于输出日志
output_box = tk.Text(root, height=15, width=600)
output_box.pack(pady=20)

# 进度条与标签
progress_label_var = tk.StringVar(value="进度：0.0%")
progress_label = tk.Label(root, textvariable=progress_label_var)
progress_label.pack(pady=2)

progress_var = tk.DoubleVar(value=0.0)
progress_bar = ttk.Progressbar(root, orient="horizontal", length=800, mode="determinate", maximum=100.0, variable=progress_var)
progress_bar.pack(pady=6)

button2 = tk.Button(root, text="处理Excel", command=button2_click)
button2.pack(pady=10)

button1 = tk.Button(root, text="处理TXT", command=button1_click)
button1.pack(pady=10)

root.mainloop()

# 能不能再让软件输出一版没有“【】”里面的内容、格式与源文件一致、而且前面一样的还是标黄的。然后再加一个统计功能，源文件里这一段一共多少句（两个标点符号之间的就算一句，无论是逗号还是句号等等），存在标黄的有多少句，都标出来。